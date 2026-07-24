#!/usr/bin/env python3
"""Fill the approved weekly-report DOCX template without rebuilding its layout."""

from __future__ import annotations

import argparse
import json
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt


FONT_NAME = "바탕"


def clean_prefix(value: object, prefixes: tuple[str, ...]) -> str:
    text = " ".join(str(value or "").split()).strip()
    for prefix in prefixes:
        if text.startswith(prefix):
            return text[len(prefix) :].strip()
    return text


def style_run(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def find_paragraph(doc: Document, marker: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            return paragraph
    raise ValueError(f"Template marker not found: {marker}")


def replace_token_in_paragraph(paragraph: Paragraph, token: str, value: str) -> None:
    if token not in paragraph.text:
        return
    for run in paragraph.runs:
        if token in run.text:
            run.text = run.text.replace(token, value)
            return
    original = paragraph.runs[0] if paragraph.runs else None
    clear_runs(paragraph)
    run = paragraph.add_run(paragraph.text.replace(token, value))
    if original is not None:
        style_run(run, original.font.size.pt if original.font.size else 14, bold=bool(original.bold))


def clone_marker_before(marker: Paragraph) -> Paragraph:
    clone = deepcopy(marker._p)
    marker._p.addprevious(clone)
    return Paragraph(clone, marker._parent)


def write_report_item(paragraph: Paragraph, item: dict) -> None:
    clear_runs(paragraph)
    main = clean_prefix(item.get("main"), ("- ", "• "))
    if not main:
        main = "특이사항 없음"
    run = paragraph.add_run(main)
    style_run(run, 14)
    for sub in item.get("subs") or []:
        run = paragraph.add_run()
        run.add_break()
        run.text = f"* {clean_prefix(sub, ('* ', '- ', '• '))}"
        style_run(run, 13)
    implication = clean_prefix(item.get("implication"), ("☞ ", "- ", "• "))
    if implication:
        run = paragraph.add_run()
        run.add_break()
        run.text = f"☞ {implication}"
        style_run(run, 13)


def fill_repeating_items(doc: Document, marker_text: str, items: list[dict]) -> None:
    marker = find_paragraph(doc, marker_text)
    safe_items = items or [{"main": "특이사항 없음", "subs": [], "implication": ""}]
    for item in safe_items:
        paragraph = clone_marker_before(marker)
        write_report_item(paragraph, item)
    remove_paragraph(marker)


def fill_impacts(doc: Document, impacts: list[str]) -> None:
    marker = find_paragraph(doc, "{{IMPACT_ITEMS}}")
    safe = impacts or ["주요 정세 변화에 따른 현장 안전 및 사업 영향 지속 점검"]
    for value in safe[:2]:
        paragraph = clone_marker_before(marker)
        clear_runs(paragraph)
        run = paragraph.add_run(clean_prefix(value, ("• ", "- ")))
        style_run(run, 14, bold=True)
    remove_paragraph(marker)


def set_cell_value(cell, value: object) -> None:
    paragraph = cell.paragraphs[0]
    size = paragraph.runs[0].font.size.pt if paragraph.runs and paragraph.runs[0].font.size else 11
    bold = bool(paragraph.runs[0].bold) if paragraph.runs else False
    clear_runs(paragraph)
    run = paragraph.add_run(str(value if value not in (None, "") else "-"))
    style_run(run, size, bold=bold)


def fill_tables(doc: Document, data: dict) -> None:
    terror = data.get("terrorStats") or {}
    terror_values = [
        "건수",
        terror.get("total", 0),
        terror.get("armed_attack", 0),
        terror.get("ied", 0),
        terror.get("assassination", 0),
        terror.get("protest", 0),
        terror.get("shooting", 0),
        terror.get("suicide_bombing", 0),
    ]
    for cell, value in zip(doc.tables[0].rows[1].cells, terror_values):
        set_cell_value(cell, value)

    rows = list(data.get("oilRows") or [])[:2]
    while len(rows) < 2:
        rows.append({"date": "-", "dubai": "-", "brent": "-", "wti": "-"})
    for table_row, item in zip(doc.tables[1].rows[1:3], rows):
        values = [item.get("date", "-"), item.get("dubai", "-"), item.get("brent", "-"), item.get("wti", "-")]
        for cell, value in zip(table_row.cells, values):
            set_cell_value(cell, value)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--data", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--latest", required=True)
    args = parser.parse_args()

    data = json.loads(Path(args.data).read_text(encoding="utf-8"))
    doc = Document(args.template)

    replace_token_in_paragraph(find_paragraph(doc, "{{PERIOD}}"), "{{PERIOD}}", str(data["period"]))
    replace_token_in_paragraph(find_paragraph(doc, "{{REPORT_DATE}}"), "{{REPORT_DATE}}", str(data["reportDate"]))
    topic = clean_prefix(data.get("internationalTopic") or "중동 주요 정세", ("• ", "- "))
    topic_paragraph = find_paragraph(doc, "{{INTERNATIONAL_TOPIC}}")
    replace_token_in_paragraph(topic_paragraph, "{{INTERNATIONAL_TOPIC}}", topic)

    sections = data.get("sections") or {}
    fill_repeating_items(doc, "{{POLITICS_ITEMS}}", list(sections.get("politics") or []))
    fill_repeating_items(doc, "{{SECURITY_ITEMS}}", list(sections.get("terror_security") or []))
    fill_repeating_items(doc, "{{ECONOMY_ITEMS}}", list(sections.get("oil_economy") or []))
    fill_repeating_items(doc, "{{INTERNATIONAL_ITEMS}}", list(sections.get("regional") or []))
    fill_impacts(doc, list(data.get("groupImpacts") or []))
    fill_tables(doc, data)

    output = Path(args.output)
    latest = Path(args.latest)
    output.parent.mkdir(parents=True, exist_ok=True)
    latest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    shutil.copyfile(output, latest)


if __name__ == "__main__":
    main()
